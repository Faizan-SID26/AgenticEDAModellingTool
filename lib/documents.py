"""Domain-document reader.

Industrial projects depend heavily on process knowledge — Process
Understanding Documents (PUDs), engineering specifications, SOPs, prior
investigations. The framework treats these as first-class priors:

- At ``/init`` we scan ``<project>/data/`` for ``.md``, ``.txt``,
  ``.pdf``, ``.docx`` and produce ``memory/DOMAIN_DOCS.md`` containing
  the extracted text plus a per-file summary header.
- At ``/plan`` the planner reads ``DOMAIN_DOCS.md`` *before* asking
  questions and incorporates it into the question batch (e.g., it
  surfaces named process stages, expected drivers, known failure modes
  rather than re-asking).

Heavy parsers (PDF, DOCX) are *optional* dependencies — the module
tolerates them being absent and records a `parser_missing` note for the
planner to surface to the user.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional

_log = logging.getLogger("eda.documents")

_SUPPORTED_SUFFIXES: tuple[str, ...] = (".md", ".txt", ".pdf", ".docx", ".rtf")
_TEXT_BUDGET_BYTES = 200_000
"""Per-file text-extraction cap. PUDs are typically <50KB of prose;
beyond this we truncate to keep the planner's context bounded."""


@dataclass
class DomainDocument:
    """One extracted document, with provenance and any parse warnings."""

    path: str
    """Project-relative path."""

    suffix: str
    """File extension (".md", ".pdf", ...)."""

    n_chars: int = 0
    """Length of extracted text (after truncation)."""

    text: str = ""
    """Extracted plain text. Empty if parsing failed."""

    error: Optional[str] = None
    """Parse error, if any (e.g., 'parser_missing:pypdf')."""

    truncated: bool = False
    """True if the file was longer than the per-file budget."""


@dataclass
class DomainCorpus:
    """Aggregate of every DomainDocument found at /init."""

    documents: list[DomainDocument] = field(default_factory=list)
    parser_warnings: list[str] = field(default_factory=list)

    def total_chars(self) -> int:
        return sum(d.n_chars for d in self.documents)

    def successful(self) -> list[DomainDocument]:
        return [d for d in self.documents if not d.error and d.text]


# --- Per-format extractors ---------------------------------------------


def _read_text_file(path: Path) -> str:
    """Read a plain-text-ish file with a robust encoding fallback."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> tuple[str, Optional[str]]:
    """Return (text, error). Tries pypdf, then pdfplumber."""
    text_parts: list[str] = []
    try:
        import pypdf  # type: ignore

        reader = pypdf.PdfReader(str(path))
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001 — never abort on one page
                continue
        return ("\n".join(text_parts).strip(), None)
    except ImportError:
        pass
    except Exception as e:  # noqa: BLE001
        _log.debug("pypdf failed for %s: %s", path, e)

    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception:  # noqa: BLE001
                    continue
        return ("\n".join(text_parts).strip(), None)
    except ImportError:
        return (
            "",
            "parser_missing:pypdf or pdfplumber — install one to read PDFs",
        )
    except Exception as e:  # noqa: BLE001
        return ("", f"pdf_parse_failed: {e}")


def _read_docx(path: Path) -> tuple[str, Optional[str]]:
    try:
        import docx  # type: ignore  # python-docx
    except ImportError:
        return ("", "parser_missing:python-docx — pip install python-docx to read DOCX")
    try:
        d = docx.Document(str(path))
        return ("\n".join(p.text for p in d.paragraphs).strip(), None)
    except Exception as e:  # noqa: BLE001
        return ("", f"docx_parse_failed: {e}")


def _read_rtf(path: Path) -> tuple[str, Optional[str]]:
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
    except ImportError:
        return (
            "",
            "parser_missing:striprtf — pip install striprtf to read RTF",
        )
    try:
        return (rtf_to_text(_read_text_file(path)).strip(), None)
    except Exception as e:  # noqa: BLE001
        return ("", f"rtf_parse_failed: {e}")


# --- Public API --------------------------------------------------------


def is_document(path: Path) -> bool:
    """Return True if `path` looks like a domain document."""
    return path.is_file() and path.suffix.lower() in _SUPPORTED_SUFFIXES


def extract(path: Path, *, max_bytes: int = _TEXT_BUDGET_BYTES) -> DomainDocument:
    """Extract text from one document. Always returns a DomainDocument."""
    rel = str(path)
    suffix = path.suffix.lower()
    text = ""
    error: Optional[str] = None

    try:
        if suffix in (".md", ".txt"):
            text = _read_text_file(path)
        elif suffix == ".pdf":
            text, error = _read_pdf(path)
        elif suffix == ".docx":
            text, error = _read_docx(path)
        elif suffix == ".rtf":
            text, error = _read_rtf(path)
        else:
            error = f"unsupported_suffix:{suffix}"
    except Exception as e:  # noqa: BLE001 — never fail /init on one bad doc
        error = f"unexpected:{type(e).__name__}: {e}"

    truncated = False
    if text and len(text.encode("utf-8")) > max_bytes:
        # Truncate by characters for simplicity; UTF-8 oversize is fine.
        text = text[:max_bytes]
        truncated = True

    return DomainDocument(
        path=rel,
        suffix=suffix,
        n_chars=len(text),
        text=text,
        error=error,
        truncated=truncated,
    )


def discover(data_dir: Path) -> list[Path]:
    """Return every supported document under `data_dir` (sorted)."""
    if not data_dir.exists():
        return []
    return sorted(p for p in data_dir.iterdir() if is_document(p))


def collect(data_dir: Path) -> DomainCorpus:
    """Discover + extract every domain document under `data_dir`."""
    corpus = DomainCorpus()
    for path in discover(data_dir):
        doc = extract(path)
        corpus.documents.append(doc)
        if doc.error and doc.error.startswith("parser_missing"):
            if doc.error not in corpus.parser_warnings:
                corpus.parser_warnings.append(doc.error)
    return corpus


# --- Rendering ---------------------------------------------------------


def render_domain_docs_md(
    corpus: DomainCorpus,
    *,
    project_dir: Optional[Path] = None,
) -> str:
    """Render the corpus as a single markdown bundle the planner can consume.

    Each document is presented as a level-2 section with a small
    metadata block + the extracted text in a fenced block.
    """
    base = project_dir.resolve() if project_dir else None
    lines: list[str] = []
    lines.append("# Domain documents\n")
    lines.append(
        "_Auto-extracted at `/init`. The planner reads this before asking "
        "questions; if anything here is wrong, just say so during `/plan`._\n"
    )
    if corpus.parser_warnings:
        lines.append("## Parser warnings\n")
        for w in corpus.parser_warnings:
            lines.append(f"- `{w}`")
        lines.append("")
    if not corpus.documents:
        lines.append("_No supported documents found under `data/`. Drop `.md`, `.txt`, `.pdf`, `.docx`, or `.rtf` files there to give the planner domain context._\n")
        return "\n".join(lines) + "\n"
    for doc in corpus.documents:
        path_disp = doc.path
        if base is not None:
            try:
                path_disp = str(Path(doc.path).resolve().relative_to(base))
            except Exception:  # noqa: BLE001
                pass
        lines.append(f"## `{path_disp}`\n")
        lines.append(f"- format: `{doc.suffix}`")
        lines.append(f"- extracted_chars: {doc.n_chars}")
        if doc.truncated:
            lines.append(
                f"- _truncated to {_TEXT_BUDGET_BYTES} bytes; full file not loaded_"
            )
        if doc.error:
            lines.append(f"- error: `{doc.error}`")
        lines.append("")
        if doc.text:
            lines.append("```text")
            lines.append(doc.text)
            lines.append("```\n")
        else:
            lines.append("_(no text extracted)_\n")
    return "\n".join(lines) + "\n"


def write_domain_docs(project_dir: Path, corpus: DomainCorpus) -> Optional[Path]:
    """Persist `memory/DOMAIN_DOCS.md`. Returns the path if anything was
    written; None if there were no documents (no file is created in
    that case so /plan can detect 'no domain docs' cleanly)."""
    if not corpus.documents:
        return None
    out = Path(project_dir) / "memory" / "DOMAIN_DOCS.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(render_domain_docs_md(corpus, project_dir=project_dir), encoding="utf-8")
    return out
